import os
import datetime
from typing import List, Tuple, Any, Dict, Optional

from core.bot import Bot
from core.data import custom_data_chunk, DataChunkBase, DataManagerError
from core.command.const import *
from core.command import UserCommandBase, custom_user_command
from core.command import BotCommandBase, BotSendMsgCommand, BotSendFileCommand
from core.communication import MessageMetaData, GroupMessagePort
from utils.time import get_current_date_str, get_current_date_raw, datetime_to_str
from utils.logger import dice_log

# DataChunk 标识 / 存储结构
DC_LOG_SESSION = "log_session"
DCK_ACTIVE = "active"          # 是否正在记录
DCK_START_TIME = "start_time"  # 记录开始时间 (str)
DCK_RECORDS = "records"        # 记录的消息列表 List[Dict]
DCK_COLOR_MAP = "color_map"    # 用户ID -> 颜色 hex
DCK_MSG_COUNT = "msg_count"    # 累计消息数
DCK_LAST_HOUR_WARN = "last_hour_warn"  # 上次小时警告时间 (str)
DCK_FILTER_OUTSIDE = "filter_outside"  # 括号场外过滤
DCK_FILTER_COMMAND = "filter_command"  # 指令过滤
DCK_FILTER_BOT = "filter_bot"          # 骰娘自身过滤
DCK_FILTER_MEDIA = "filter_media"      # 图片表情过滤
DCK_FILTER_FORUM_CODE = "filter_forum_code"  # 论坛代码生成开关 (历史名称保留兼容)


@custom_data_chunk(identifier=DC_LOG_SESSION)
class _(DataChunkBase):  # noqa: E742 (单字符类名与其他 DataChunk 保持一致写法)
    def __init__(self):
        super().__init__()


COLOR_POOL = [
    "FF0000", "1E90FF", "228B22", "FF8C00", "9400D3", "DC143C", "20B2AA", "8B4513",
    "FF1493", "2E8B57", "4169E1", "DAA520", "C71585", "008B8B", "B03060", "556B2F",
]


def _pick_color(color_map: Dict[str, str], user_id: str) -> str:
    if user_id not in color_map:
        color_map[user_id] = COLOR_POOL[len(color_map) % len(COLOR_POOL)]
    return color_map[user_id]


def append_log_record(bot: Bot, group_id: str, user_id: str, nickname: str, content: str, message_id: Optional[str] = None):
    """对外工具：向当前活动日志追加一条记录（若未开启则静默）。"""
    try:
        active = bot.data_manager.get_data(DC_LOG_SESSION, [group_id, DCK_ACTIVE], False)
        if not active:
            return
        try:
            records: List[Dict] = bot.data_manager.get_data(DC_LOG_SESSION, [group_id, DCK_RECORDS])
        except DataManagerError:
            records = []
            bot.data_manager.set_data(DC_LOG_SESSION, [group_id, DCK_RECORDS], records)
        color_map: Dict[str, str] = bot.data_manager.get_data(DC_LOG_SESSION, [group_id, DCK_COLOR_MAP], {})
        _pick_color(color_map, user_id)
        bot.data_manager.set_data(DC_LOG_SESSION, [group_id, DCK_COLOR_MAP], color_map)
        # 昵称兜底：bot消息使用默认“骰娘”，用户消息尝试从 nickname 系统获取
        if (not nickname) or nickname in ("UNDEF_NAME", "----"):
            if user_id == bot.account:
                nickname = "骰娘"
            else:
                try:
                    nickname = bot.get_nickname(user_id, group_id) or user_id
                except Exception:
                    nickname = user_id
        # 过滤（bot 写入也可能被用户要求过滤掉）
        if not should_filter_record(bot, group_id, user_id, content, is_bot=True):
            record = {
                "time": get_current_date_str(),
                "user_id": user_id,
                "nickname": nickname or user_id,
                "content": content,
            }
            if message_id:
                record["message_id"] = message_id
            records.append(record)
            bot.data_manager.set_data(DC_LOG_SESSION, [group_id, DCK_RECORDS], records)
    except Exception:
        # 不让日志记录的异常影响主流程
        pass


def _generate_forum_code_log(records: List[Dict]) -> str:
    """生成论坛代码格式的日志内容。"""
    forum_code = []
    for record in records:
        time = record.get("time", "未知时间")
        nickname = record.get("nickname", "未知用户")
        content = record.get("content", "")
        forum_code.append(f"[color=#9ca3af]{time}[/color][color=#f99252] <{nickname}>{content} [/color]")
    return "\n".join(forum_code)


def should_filter_record(bot: Bot, group_id: str, user_id: str, content: str, is_bot: bool = False) -> bool:
    """根据当前群的过滤配置判断是否过滤此消息。
    过滤规则：
    - outside: 消息整体被括号包裹 (支持全角/半角 ()（）) 时过滤
    - command: 以 '.' 或 '。' 开头的指令过滤
    - bot: 机器人自身消息过滤 (is_bot=True 时)
    - media: 图片/表情 CQ 码过滤
    - 文件 CQ 永远不过滤逻辑直接剔除

    forum_code(原“论坛代码过滤”) 已改为“论坛代码生成”，只控制是否额外生成带论坛色彩标签的 txt 文件，不再影响过滤逻辑。
    """
    try:
        outside_f = bot.data_manager.get_data(DC_LOG_SESSION, [group_id, DCK_FILTER_OUTSIDE], False)
        command_f = bot.data_manager.get_data(DC_LOG_SESSION, [group_id, DCK_FILTER_COMMAND], False)
        bot_f = bot.data_manager.get_data(DC_LOG_SESSION, [group_id, DCK_FILTER_BOT], False)
    except Exception:
        return False

    text = content.strip()
    # outside: 判断是否用任意一对括号包裹
    if outside_f and (
        (text.startswith('(') and text.endswith(')')) or
        (text.startswith('（') and text.endswith('）'))
    ):
        return True
    if command_f and (text.startswith('.') or text.startswith('。')):
        return True
    if bot_f and is_bot:
        return True

    # 媒体过滤（图片/表情）
    try:
        media_f = bot.data_manager.get_data(DC_LOG_SESSION, [group_id, DCK_FILTER_MEDIA], False)
    except Exception:
        media_f = False
    if media_f:
        lowered = text.lower()
        if ("[cq:image," in lowered) or ("[cq:face," in lowered) or ("[cq:emoji," in lowered):
            return True

    # 文件永远不记录
    if "[CQ:file," in content:
        return True

    return False


LOC_LOG_ON_START = "log_on_start"            # 开始记录时
LOC_LOG_ON_ALREADY = "log_on_already"        # 已在记录
LOC_LOG_OFF_NOT_ACTIVE = "log_off_not_active"  # 关闭但未开启
LOC_LOG_OFF_RESULT = "log_off_result"        # 关闭并生成，{count}
LOC_LOG_USAGE = "log_usage"                  # 用法说明
LOC_LOG_SET_MENU = "log_set_menu"            # 设置菜单
LOC_LOG_SET_TOGGLED = "log_set_toggled"      # 切换结果
LOC_LOG_HELP = "log_help"                    # 完整帮助
LOC_LOG_FOLDER_FAIL = "log_folder_fail"      # 创建群文件夹失败
LOC_LOG_FOLDER_HINT = "log_folder_hint"      # 创建群文件夹提示
LOC_LOG_FOLDER_POLICY = "log_folder_policy"  # 启动时提示上传策略


@custom_user_command(readable_name="跑团日志指令", priority=DPP_COMMAND_PRIORITY_DEFAULT,
                     flag=DPP_COMMAND_FLAG_DEFAULT, cluster=DPP_COMMAND_CLUSTER_DEFAULT, group_only=True)
class LogCommand(UserCommandBase):
    """
    .log on  开始记录本群消息
    .log off 停止并生成日志文件
    """

    def __init__(self, bot: Bot):
        super().__init__(bot)
        # 移除本地化注册，直接写死文本内容
        self.log_on_start = "新的故事开始了，祝旅途愉快！\n记录已经开启。"
        self.log_on_already = "日志记录已在进行中。"
        self.log_off_not_active = "当前没有进行中的日志记录。使用 .log on 开始。"
        self.log_off_result = "日志已生成，共 {count} 条消息。"
        self.log_usage = "用法: .log on 开始记录；.log off 结束并生成。"
        self.log_set_menu = (
            "日志过滤/输出设置:\n"
            "1： \noutside / 场外发言过滤\n启用： {outside} \n"
            "2： \ncommand / 指令过滤\n启用： {command} \n"
            "3： \nbot / bot过滤\n启用： {bot} \n"
            "4： \nmedia / 图片表情过滤\n启用： {media} \n"
            "5： \nforum_code / 论坛代码生成\n启用： {forum_code} \n"
            "切换示例: .log set outside  或  .log set 论坛代码生成\n"
            "仅输入 .log set 显示本菜单。"
        )
        self.log_help = (
            "日志功能指令: \n"
            ".log on  开始记录\n"
            ".log off 结束记录并生成文件(docx+txt)\n"
            ".log set  查看/切换过滤开关 (支持: outside/场外发言过滤, command/指令过滤, bot/bot过滤, media/图片表情过滤, forum_code/论坛代码生成)\n"
            ".log set <选项> 切换对应过滤。\n"
            "过滤说明: \n场外发言过滤=过滤括号内信息\n指令过滤=过滤以句号起始指令\nbot过滤=过滤骰娘自身消息\n图片表情过滤=过滤图片和表情消息。"
        )
        self.log_folder_fail = "权限不足，无法创建跑团log文件夹，已上传至根目录。"
        self.log_folder_hint = "日志文件将上传到群文件夹: 跑团log"
        self.log_folder_policy = "说明：若群文件存在‘跑团log’文件夹，日志文件将上传到该文件夹；若不存在则上传至群文件根目录。"
        self.bot.loc_helper.register_loc_text(LOC_LOG_SET_TOGGLED, "{item} 切换为 {state}", "切换日志设置")

    def can_process_msg(self, msg_str: str, meta: MessageMetaData) -> Tuple[bool, bool, Any]:
        should_proc = msg_str.startswith(".log")
        if not should_proc:
            return False, False, None
        args = msg_str.split()
        action = args[1].lower() if len(args) > 1 else ""
        param = args[2].lower() if len(args) > 2 else ""
        return True, False, (action, param)  # 不继续传递

    def process_msg(self, msg_str: str, meta: MessageMetaData, hint: Any) -> List[BotCommandBase]:
        action, param = hint if isinstance(hint, tuple) else (hint, "")
        group_id = meta.group_id
        if not group_id:
            return []

        feedback = ""
        active = self.bot.data_manager.get_data(DC_LOG_SESSION, [group_id, DCK_ACTIVE], False)
        if action == "on":
            if active:
                feedback = self.log_on_already
            else:
                start_time = get_current_date_str()
                self.bot.data_manager.set_data(DC_LOG_SESSION, [group_id, DCK_ACTIVE], True)
                self.bot.data_manager.set_data(DC_LOG_SESSION, [group_id, DCK_START_TIME], start_time)
                self.bot.data_manager.set_data(DC_LOG_SESSION, [group_id, DCK_RECORDS], [])
                self.bot.data_manager.set_data(DC_LOG_SESSION, [group_id, DCK_COLOR_MAP], {})
                self.bot.data_manager.set_data(DC_LOG_SESSION, [group_id, DCK_MSG_COUNT], 0)
                self.bot.data_manager.set_data(DC_LOG_SESSION, [group_id, DCK_LAST_HOUR_WARN], start_time)
                # 只在缺失时初始化过滤配置，避免覆盖用户此前通过 .log set 设定的开关
                for flag in (DCK_FILTER_OUTSIDE, DCK_FILTER_COMMAND, DCK_FILTER_BOT, DCK_FILTER_MEDIA, DCK_FILTER_FORUM_CODE):
                    try:
                        _ = self.bot.data_manager.get_data(DC_LOG_SESSION, [group_id, flag])
                    except DataManagerError:
                        self.bot.data_manager.set_data(DC_LOG_SESSION, [group_id, flag], False)
                feedback = self.log_on_start + "\n" + self.log_folder_policy
        elif action == "off":
            if not active:
                feedback = self.log_off_not_active
            else:
                # 结束 & 生成文件
                records = self.bot.data_manager.get_data(DC_LOG_SESSION, [group_id, DCK_RECORDS], [])
                start_time = self.bot.data_manager.get_data(DC_LOG_SESSION, [group_id, DCK_START_TIME], "unknown_start")
                color_map = self.bot.data_manager.get_data(DC_LOG_SESSION, [group_id, DCK_COLOR_MAP], {})
                self.bot.data_manager.set_data(DC_LOG_SESSION, [group_id, DCK_ACTIVE], False)
                file_main_path, display_name, extra_files = self._generate_file(group_id, start_time, records, color_map)
                feedback = self.log_off_result.format(count=len(records))
                port = GroupMessagePort(group_id)
                folder_prefix = "跑团log/"  # 仅当群里已存在该文件夹时适配器会放进去；否则上传到根目录
                cmds: List[BotCommandBase] = [BotSendMsgCommand(self.bot.account, feedback, [port]),
                                              BotSendFileCommand(self.bot.account, file_main_path, folder_prefix + display_name, [port])]
                for fpath, fname in extra_files:
                    cmds.append(BotSendFileCommand(self.bot.account, fpath, folder_prefix + fname, [port]))
                return cmds
        elif action == "set":
            # 允许未开启时预配置
            for base_key in (DCK_RECORDS, DCK_COLOR_MAP):
                try:
                    _ = self.bot.data_manager.get_data(DC_LOG_SESSION, [group_id, base_key])
                except DataManagerError:
                    self.bot.data_manager.set_data(DC_LOG_SESSION, [group_id, base_key], [] if base_key == DCK_RECORDS else {})
            for flag in (DCK_FILTER_OUTSIDE, DCK_FILTER_COMMAND, DCK_FILTER_BOT, DCK_FILTER_MEDIA, DCK_FILTER_FORUM_CODE):
                try:
                    _ = self.bot.data_manager.get_data(DC_LOG_SESSION, [group_id, flag])
                except DataManagerError:
                    self.bot.data_manager.set_data(DC_LOG_SESSION, [group_id, flag], False)
            if not param:
                outside = self.bot.data_manager.get_data(DC_LOG_SESSION, [group_id, DCK_FILTER_OUTSIDE], False)
                command_f = self.bot.data_manager.get_data(DC_LOG_SESSION, [group_id, DCK_FILTER_COMMAND], False)
                bot_f = self.bot.data_manager.get_data(DC_LOG_SESSION, [group_id, DCK_FILTER_BOT], False)
                media_f = self.bot.data_manager.get_data(DC_LOG_SESSION, [group_id, DCK_FILTER_MEDIA], False)
                forum_code_f = self.bot.data_manager.get_data(DC_LOG_SESSION, [group_id, DCK_FILTER_FORUM_CODE], False)
                feedback = self.log_set_menu.format(
                                           outside="ON" if outside else "OFF",
                                           command="ON" if command_f else "OFF",
                                           bot="ON" if bot_f else "OFF",
                                           media="ON" if media_f else "OFF",
                                           forum_code="ON" if forum_code_f else "OFF") + (" (记录未开启)" if not active else "")
            else:
                key_map = {
                    "outside": DCK_FILTER_OUTSIDE,
                    "场外发言过滤": DCK_FILTER_OUTSIDE,
                    "command": DCK_FILTER_COMMAND,
                    "指令过滤": DCK_FILTER_COMMAND,
                    "bot": DCK_FILTER_BOT,
                    "bot过滤": DCK_FILTER_BOT,
                    "media": DCK_FILTER_MEDIA,
                    "图片表情过滤": DCK_FILTER_MEDIA,
                    "forum_code": DCK_FILTER_FORUM_CODE,
                    "论坛代码过滤": DCK_FILTER_FORUM_CODE,  # 兼容旧说法
                    "论坛代码生成": DCK_FILTER_FORUM_CODE,
                    "论坛": DCK_FILTER_FORUM_CODE,
                    "论坛代码": DCK_FILTER_FORUM_CODE,
                }
                if param in key_map:
                    cur = self.bot.data_manager.get_data(DC_LOG_SESSION, [group_id, key_map[param]], False)
                    self.bot.data_manager.set_data(DC_LOG_SESSION, [group_id, key_map[param]], not cur)
                    dice_log(f"[LogCommand] 切换过滤设置: {param} -> {'ON' if not cur else 'OFF'}")
                    feedback = self.format_loc(LOC_LOG_SET_TOGGLED, item=param, state="ON" if (not cur) else "OFF") + (" (记录未开启，已保存)" if not active else "")
                else:
                    feedback = "未知选项，可用: outside/场外发言过滤 | command/指令过滤 | bot/bot过滤 | media/图片表情过滤 | forum_code/论坛代码过滤"
        elif action == "":
            # 仅 .log -> 完整帮助
            feedback = self.log_help
        else:
            feedback = self.log_usage

        return [BotSendMsgCommand(self.bot.account, feedback, [GroupMessagePort(group_id)])]

    def _generate_file(self, group_id: str, start_time: str, records: List[Dict], color_map: Dict[str, str]) -> Tuple[str, str, List[Tuple[str,str]]]:
        """生成 docx 和 txt 文件。
        变更：
        1. 不再使用 '---- 时间' 折叠同一用户的连续发言，避免用户误解昵称丢失。
        2. 将常见 CQ 码（at/reply/image/file）转换为更易读的文本。
        3. 若为回复（reply）尝试根据 message_id 在已记录消息中找到原文前若干字符作为引用预览。
        返回: (主文件路径, 主文件名, 额外文件列表[(path,name),...])"""
        logs_dir = os.path.join(self.bot.data_path, "logs")
        os.makedirs(logs_dir, exist_ok=True)
        # start_time 形如 YYYY/MM/DD HH:MM:SS，需要去掉 / : 等不安全字符
        safe_start = (start_time
                      .replace('/', '-')
                      .replace(':', '-')
                      .replace(' ', '_'))
        display_name_base = f"log_{group_id}_{safe_start}"
        # 当前群 .nn 生效昵称（优先）
        import re as _re_outer
        nn_cache: Dict[str,str] = {}
        for r in records:
            uid = r.get('user_id')
            if uid and uid not in nn_cache:
                try:
                    val = self.bot.get_nickname(uid, group_id)
                except Exception:
                    val = None
                if val and val not in ("UNDEF_NAME", "----"):
                    nn_cache[uid] = val
        # 构建 message_id -> {'content':..., 'nickname':...}
        msg_map: Dict[str, Dict[str, str]] = {}
        for r in records:
            mid = r.get('message_id')
            if not mid:
                continue
            raw_content = r.get('content', '')
            cleaned = _re_outer.sub(r"\[CQ:reply,(?:id|reply|source_id)=\d+[^\]]*\]", "", raw_content)
            display_nick = nn_cache.get(r.get('user_id')) or r.get('nickname','?')
            msg_map[str(mid)] = {'content': cleaned, 'nickname': display_nick}

        # 预构建 user_id -> 昵称 映射 (at 展示用)
        user_nick: Dict[str, str] = {}
        for r in records:
            uid = r.get('user_id')
            if uid and uid not in user_nick:
                user_nick[uid] = nn_cache.get(uid) or r.get('nickname') or uid

        def humanize_cq(raw: str) -> str:
            import re
            out = raw
            # 先处理 reply
            def repl_reply(m):
                rid = m.group(1)
                origin = msg_map.get(rid)
                if not origin:
                    return "| 引用消息不在 log 范围内\n"
                origin_content = origin['content'].strip() or "(空白)"
                lines = [ln.strip() for ln in origin_content.splitlines() if ln.strip()][:3] or [origin_content]
                lines = [ln[:60] + ('…' if len(ln) > 60 else '') for ln in lines]
                quote_lines = [f"| {origin['nickname']}"] + [f"| {ln}" for ln in lines]
                return "\n".join(quote_lines) + "\n"
            try:
                out = re.sub(r"\[CQ:reply,(?:id|reply|source_id)=(\d+)[^\]]*\]", repl_reply, out)
            except Exception:
                pass

            # 再处理 @
            def repl_at(m):
                uid = m.group(1)
                nick = user_nick.get(uid)
                if not nick or nick in ("UNDEF_NAME", "----"):
                    try:
                        nick = self.bot.get_nickname(uid, group_id) or uid
                    except Exception:
                        nick = uid
                return f"@{nick}"
            try:
                out = re.sub(r"\[CQ:at,qq=(\d+)(?:,[^\]]*)?\]", repl_at, out)
            except Exception:
                pass

            # 图片/表情 CQ 原样保留；文件在记录阶段已过滤
            return out

        txt_path = os.path.join(logs_dir, display_name_base + ".txt")
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write(f"群 {group_id} 跑团日志 (开始于 {start_time})\n\n")
            for rec in records:
                # 修正机器人昵称
                if rec.get('user_id') == self.bot.account and (not rec.get('nickname') or rec.get('nickname') in ("UNDEF_NAME", "----")):
                    rec['nickname'] = "骰娘"
                disp = nn_cache.get(rec.get('user_id')) or rec.get('nickname','?')
                content_out = humanize_cq(rec.get('content', ''))
                f.write(f"{disp} ({rec.get('user_id','?')})  {rec.get('time','?')}\n")
                f.write(content_out + "\n\n")

        docx_path = None
        try:
            from docx import Document  # type: ignore
            from docx.shared import RGBColor  # type: ignore
            doc = Document()
            doc.add_heading(f"群 {group_id} 跑团日志 (开始于 {start_time})", level=1)
            for rec in records:
                uid = rec.get('user_id','?')
                color_hex = color_map.get(uid, "000000")
                r = int(color_hex[0:2], 16); g = int(color_hex[2:4], 16); b = int(color_hex[4:6], 16)
                if rec.get('user_id') == self.bot.account and (not rec.get('nickname') or rec.get('nickname') in ("UNDEF_NAME", "----")):
                    rec['nickname'] = "骰娘"
                disp = nn_cache.get(uid) or rec.get('nickname','?')
                header = doc.add_paragraph()
                run1 = header.add_run(f"{disp} ({uid})  {rec.get('time','?')}")
                run1.font.color.rgb = RGBColor(r, g, b)
                body_p = doc.add_paragraph()
                run2 = body_p.add_run(humanize_cq(rec.get('content','')))
                run2.font.color.rgb = RGBColor(r, g, b)
            docx_path = os.path.join(logs_dir, display_name_base + ".docx")
            doc.save(docx_path)
        except Exception as e:
            # 记录失败原因，便于诊断（如服务器环境丢失依赖、权限问题等）
            try:
                dice_log(f"[LogExport] docx generation failed: {type(e).__name__}: {e}")
            except Exception:
                pass
            docx_path = None

        extra: List[Tuple[str,str]] = []
        # forum_code 额外论坛格式文件
        try:
            forum_code_on = self.bot.data_manager.get_data(DC_LOG_SESSION, [group_id, DCK_FILTER_FORUM_CODE], False)
        except Exception:
            forum_code_on = False
        if forum_code_on:
            try:
                forum_txt_path = os.path.join(logs_dir, display_name_base + "_forum.txt")
                with open(forum_txt_path, "w", encoding="utf-8") as ff:
                    ff.write(_generate_forum_code_log(records))
                # 若 docx 为主文件，则 forum 与普通 txt 都作为 extra；若 txt 为主文件则只追加 forum
                extra.append((forum_txt_path, os.path.basename(forum_txt_path)))
            except Exception as e:
                try:
                    dice_log(f"[LogExport] forum code generation failed: {type(e).__name__}: {e}")
                except Exception:
                    pass
        if docx_path:
            # 主文件用docx，额外上传txt
            extra.append((txt_path, os.path.basename(txt_path)))
            return docx_path, os.path.basename(docx_path), extra
        else:
            return txt_path, os.path.basename(txt_path), extra

    def get_help(self, keyword: str, meta: MessageMetaData) -> str:
        if keyword in ("log", "日志"):
            return self.log_usage
        return ""

    def get_description(self) -> str:
        return ".log on/off 记录跑团日志"


@custom_user_command(readable_name="跑团日志记录器", priority=DPP_COMMAND_PRIORITY_DEFAULT - 50,
                     flag=0, cluster=DPP_COMMAND_CLUSTER_DEFAULT, group_only=True)
class LogRecorderCommand(UserCommandBase):
    """负责在开启状态下采集所有普通消息 (不以'.'开头)"""

    def __init__(self, bot: Bot):
        super().__init__(bot)

    def can_process_msg(self, msg_str: str, meta: MessageMetaData) -> Tuple[bool, bool, Any]:
        if not meta.group_id:
            return False, False, None
        is_active = self.bot.data_manager.get_data(DC_LOG_SESSION, [meta.group_id, DCK_ACTIVE], False)
        if not is_active:
            return False, False, None
        # 现在指令也记录；继续传递给后续命令
        return True, True, None

    def process_msg(self, msg_str: str, meta: MessageMetaData, hint: Any) -> List[BotCommandBase]:
        group_id = meta.group_id
        cmds: List[BotCommandBase] = []
        # 读取当前数据
        try:
            records: List[Dict] = self.bot.data_manager.get_data(DC_LOG_SESSION, [group_id, DCK_RECORDS])
        except DataManagerError:
            records = []
            self.bot.data_manager.set_data(DC_LOG_SESSION, [group_id, DCK_RECORDS], records)
        color_map: Dict[str, str] = self.bot.data_manager.get_data(DC_LOG_SESSION, [group_id, DCK_COLOR_MAP], {})
        _pick_color(color_map, meta.user_id)
        self.bot.data_manager.set_data(DC_LOG_SESSION, [group_id, DCK_COLOR_MAP], color_map)

        # 内容策略：判断原始 raw 是否以半角 . 或全角 。 开头，原样记录；否则记录 raw（保留CQ码）
        raw = getattr(meta, 'raw_msg', '') or msg_str
        if raw.startswith('.') or raw.startswith('。'):
            content = raw  # 指令保持原样
        else:
            content = raw

        # 过滤逻辑
        if should_filter_record(self.bot, group_id, meta.user_id, content):
            return []

        # 计数器与时间提醒逻辑
        try:
            msg_count = self.bot.data_manager.get_data(DC_LOG_SESSION, [group_id, DCK_MSG_COUNT], 0)
        except Exception:
            msg_count = 0
        msg_count += 1
        self.bot.data_manager.set_data(DC_LOG_SESSION, [group_id, DCK_MSG_COUNT], msg_count)

        # 每100条提醒
        if msg_count % 100 == 0:
            cmds.append(BotSendMsgCommand(self.bot.account, f"当前log消息已有{msg_count}条。", [GroupMessagePort(group_id)]))

        # 每小时提醒
        import datetime
        try:
            start_time_str = self.bot.data_manager.get_data(DC_LOG_SESSION, [group_id, DCK_START_TIME], get_current_date_str())
            last_warn_str = self.bot.data_manager.get_data(DC_LOG_SESSION, [group_id, DCK_LAST_HOUR_WARN], start_time_str)
        except Exception:
            start_time_str = get_current_date_str()
            last_warn_str = start_time_str
        try:
            start_dt = datetime.datetime.strptime(start_time_str, "%Y/%m/%d %H:%M:%S")
            last_warn_dt = datetime.datetime.strptime(last_warn_str, "%Y/%m/%d %H:%M:%S")
            now_dt = datetime.datetime.now()
            hours_enabled = int((now_dt - start_dt).total_seconds() // 3600)
            hours_last_warn = int((last_warn_dt - start_dt).total_seconds() // 3600)
            if hours_enabled > hours_last_warn:
                cmds.append(BotSendMsgCommand(self.bot.account, f"当前log已启用{hours_enabled}小时，如果已经save请注意关闭log。", [GroupMessagePort(group_id)]))
                self.bot.data_manager.set_data(DC_LOG_SESSION, [group_id, DCK_LAST_HOUR_WARN], now_dt.strftime("%Y/%m/%d %H:%M:%S"))
        except Exception:
            pass

        # 真正写入log
        records.append({
            "time": get_current_date_str(),
            "user_id": meta.user_id,
            "nickname": meta.nickname or meta.user_id,
            "content": content,
        })
        self.bot.data_manager.set_data(DC_LOG_SESSION, [group_id, DCK_RECORDS], records)
        return cmds

    def get_help(self, keyword: str, meta: MessageMetaData) -> str:
        return ""

    def get_description(self) -> str:
        return ""  # 隐藏

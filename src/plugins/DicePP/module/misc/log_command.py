"""
跑团记录指令
"""

from typing import List, Tuple, Any
import datetime
import os
import smtplib
import docx

from core.bot import Bot
from core.data import DataChunkBase, custom_data_chunk
from core.config import DATA_PATH, CFG_MASTER, CFG_ADMIN
from core.command.const import *
from core.command import UserCommandBase, custom_user_command
from core.command import BotCommandBase, BotSendMsgCommand, BotSendFileCommand
from core.communication import MessageMetaData, PrivateMessagePort, GroupMessagePort
from core import localization

LOC_LOG_START = "log_start"
LOC_LOG_NOW = "log_now"
LOC_LOG_STOP = "log_stop"
LOC_LOG_RESUME = "log_resume"
LOC_LOG_END = "log_end"
LOC_LOG_RESTART = "log_restart"
LOC_LOG_REGEN = "log_regen"
LOC_LOG_LIST = "log_list"
LOC_LOG_NOT_FOUND = "log_not_found"
LOC_LOG_NO_RESULT = "log_no_result"
LOC_LOG_INVALID_NAME = "log_invalid_name"
LOC_LOG_ALREADY_START = "log_already_start"
LOC_LOG_NOT_START = "log_not_start"
LOC_LOG_WORD_NUMBER_HINT = "log_word_number_hint"
LOC_LOG_EMAIL = "log_email"

COLOR_LIST = [[0xF9,0x92,0x52],[0xF4,0x8C,0xB6],[0x92,0x78,0xB9],[0x3E,0x80,0xCC],[0x84,0xA5,0x9D],[0x5B,0x5E,0x71],[0xCB,0x4D,0x68]]
BIG_N = "\n          "

DC_GROUPLOG = "group_log"

WORD_NUMBER_HINT_BY = 250  # 记录多少条内容提醒一次

# 存放群设置数据
@custom_data_chunk(identifier=DC_GROUPLOG)
class _(DataChunkBase):
    def __init__(self):
        super().__init__()

#一条记录
class LogRecord:
    def __init__(self, name: str, group_id: int):
        """记录一个开启的跑团"""
        self.name = name  # 文件名
        self.file_name = DATA_PATH + "\\LogData\\" + str(group_id) + "\\" + name + ".txt"  # 文件路径
        self.colored_file_name = DATA_PATH + "\\LogData\\" + str(group_id) + "\\" + name + ".docx"  # 润色文件路径
        self.mode = 0  # 0开启状态 1暂停状态 2结束状态/错误状态
        self.group_id = group_id  # Log群号
        if not os.path.exists(DATA_PATH + "\\LogData"):
            os.mkdir(DATA_PATH + "\\LogData")
        if not os.path.exists(DATA_PATH + "\\LogData\\" + str(group_id)):
            os.mkdir(DATA_PATH + "\\LogData\\" + str(group_id))
        if not os.path.exists(self.file_name):
            open(self.file_name,"w",encoding='utf-8')

    def save_message(self, character_name:str, character_qq: int, time: datetime.datetime, message: str):
        """保存记录的文本"""
        if len(message) > 0:
            time_str: str = time.strftime("%H:%M:%S")
            with open(self.file_name,"a",encoding='utf-8') as file:
                file.writelines(character_name + "(" + str(character_qq) + ") " + time_str + "\n")
                message_list : List[str] = message.splitlines()
                for message_str in message_list:
                    file.writelines(message_str + "\n")

    def generate_docx(self):
        document = docx.Document()
        document.styles['Normal'].font.name = u'宋体'
        document.styles['Normal'].font.color.rgb = docx.shared.RGBColor(0xBB,0xBB,0xBB)
        character_color: dict[int,list] = {}
        speaking_character_name: str = ""
        speaking_character_time: str = ""
        speaking_character_qq: str = ""
        speaking_character_words: list[str] = []
        paragraph = document.add_paragraph()
        paragraph.line_spacing = docx.shared.Pt(18)
        with open(self.file_name,"r",encoding='utf-8') as file:
            for line in file:
                #先依靠时间以及括号来判定是角色的那一行还是说话的文本
                if len(line) > 12 and line[-4] == ":" and line[-7] == ":" and line[-10] == " " and line[-11] == ")" and "(" in line:
                    if speaking_character_qq and len(speaking_character_words) > 0 and len(speaking_character_words[0]) > 0:
                        start_word = speaking_character_words[0][0]  #启始符
                        if not (start_word in "(（"):
                            run = paragraph.add_run(speaking_character_time)
                            run = paragraph.add_run("<" + speaking_character_name + ">:" + (BIG_N if len(speaking_character_words) > 1 else "") + BIG_N.join(speaking_character_words) + "\n")
                            #run.font.name=u'宋体'
                            #run._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'),'宋体')
                            run.font.color.rgb = docx.shared.RGBColor(character_color[speaking_character_qq][0], character_color[speaking_character_qq][1], character_color[speaking_character_qq][2])
                    data = line.split("(")
                    data[-1] = data[-1].split(")")[0]
                    data.append(line[-9:-1])
                    speaking_character_name = data[0]
                    try:
                        speaking_character_qq = int(data[1])
                    except ValueError:
                        speaking_character_name = "???"
                        speaking_character_qq = 0
                    speaking_character_time = data[2]
                    speaking_character_words = []
                    if not speaking_character_qq in character_color:
                        character_color[speaking_character_qq] = COLOR_LIST[len(character_color) % len(COLOR_LIST)]
                else:  # if not (line.startswith("(") or line.startswith("（")):
                    speaking_character_words.append(line.strip())
            if speaking_character_qq and len(speaking_character_words) > 0 and len(speaking_character_words[0]) > 0:
                start_word = speaking_character_words[0][0]  #启始符
                if not (start_word in "(（"):
                    run = paragraph.add_run(speaking_character_time)
                    run = paragraph.add_run("<" + speaking_character_name + ">:" + (BIG_N if len(speaking_character_words) > 1 else "") + BIG_N.join(speaking_character_words))
                    #run.font.name=u'宋体'
                    run.font.color.rgb = docx.shared.RGBColor(character_color[speaking_character_qq][0], character_color[speaking_character_qq][1], character_color[speaking_character_qq][2])    
        document.save(self.colored_file_name)

        #return "本次记录已生成+润色\n请使用.log download " + self.name + "下载你的跑团记录"
        return "本次记录已生成+润色\n但目前文件系统暂时有故障，请询问骰主下载你的LOG。"

@custom_user_command(readable_name="跑团记录指令",
                     priority=-10,
                     flag=DPP_COMMAND_FLAG_FUN | DPP_COMMAND_FLAG_MANAGE)
class LogCommand(UserCommandBase):
    """
    .log指令 记录跑团的功能
    """

    def __init__(self, bot: Bot):
        super().__init__(bot)
        self.record_dict: Dict[MessagePort, LogRecord] = {}
        bot.loc_helper.register_loc_text(LOC_LOG_START, "已开始跑团记录 - {name}.txt", ".log+名称 返回的内容\n name是文件名，如果不输入将为 群名+不重名的数字")
        bot.loc_helper.register_loc_text(LOC_LOG_NOW, "当前正在记录跑团 - {name}.txt", "记录中使用.log返回的内容\n name是当前记录的文件名")
        bot.loc_helper.register_loc_text(LOC_LOG_STOP, "已暂停跑团记录 - {name}.txt", "记录中使用.log stop/off暂停跑团记录\n name是当前记录的文件名")
        bot.loc_helper.register_loc_text(LOC_LOG_RESUME, "已继续跑团记录 - {name}.txt", "记录中使用.log resume/on继续跑团记录\n name是文件名")
        bot.loc_helper.register_loc_text(LOC_LOG_END, "已结束跑团记录 - {name}.txt", "记录中使用.log end结束跑团记录\n name是文件名")
        bot.loc_helper.register_loc_text(LOC_LOG_RESTART, "已重启跑团记录 - {name}.txt", ".log restart重启跑团记录\n name是文件名，如果不输入将寻找最迟的的本群不重名文件")
        bot.loc_helper.register_loc_text(LOC_LOG_REGEN, "已重新润色跑团记录 - {name}.txt", ".log regen重润色跑团记录\n name是文件名，如果不输入将寻找最迟的的本群不重名文件")
        bot.loc_helper.register_loc_text(LOC_LOG_LIST, "本群有以下记录：\n{result}", ".log list返回的内容")
        bot.loc_helper.register_loc_text(LOC_LOG_NOT_FOUND, "没有找到{name}.txt这个记录", ".log指令内任何未找到指定跑团记录的情况下返回的内容\n name是文件名")
        bot.loc_helper.register_loc_text(LOC_LOG_NO_RESULT, "本群无记录", ".log指令内任何未找到任何跑团记录的情况下返回的内容")
        bot.loc_helper.register_loc_text(LOC_LOG_INVALID_NAME, "记录名非法", ".log得到了非法名称时的情况下返回的内容")
        bot.loc_helper.register_loc_text(LOC_LOG_ALREADY_START, "已经在记录{name}.txt了", ".log在已开启状态下再次开启时返回的内容\n name是文件名")
        bot.loc_helper.register_loc_text(LOC_LOG_NOT_START, "没有记录", ".log指令没有记录的情况下返回的内容")
        bot.loc_helper.register_loc_text(LOC_LOG_WORD_NUMBER_HINT, "({name}.txt已记录{number}条内容了)", ".log每记录N条讯息后进行的提示\n name是文件名")
        bot.loc_helper.register_loc_text(LOC_LOG_EMAIL, "这是你的跑团记录 - {name}.txt\n与润色版本 - {name}.docx", ".log结束后发送的邮件文本\n name是文件名")

    def can_process_msg(self, msg_str: str, meta: MessageMetaData) -> Tuple[bool, bool, Any]:
        if not meta.group_id:
            return False, False, None
        port = GroupMessagePort(meta.group_id)
        should_proc: bool = False
        should_pass: bool = False
        
        
        
        opened_log = self.bot.data_manager.get_data(DC_GROUPLOG, [meta.group_id,"log"],"")
        if opened_log and port not in self.record_dict:
            self.record_dict[port] = LogRecord(opened_log,meta.group_id)
            self.record_dict[port].mode = 0 if self.bot.data_manager.get_data(DC_GROUPLOG, [meta.group_id,"status"],False) else 1
        if port in self.record_dict:
            record = self.record_dict[port]
            if record.mode == 0 and not (meta.plain_msg.startswith(".") or meta.plain_msg.startswith("。")):
                self.record_dict[port].save_message(self.bot.get_nickname(meta.user_id, meta.group_id),meta.user_id,datetime.datetime.now(),meta.plain_msg)
                texts_number = self.bot.data_manager.get_data(DC_GROUPLOG, [meta.group_id,"texts"],0) + 1
                self.bot.data_manager.set_data(DC_GROUPLOG, [meta.group_id,"texts"],texts_number)
                if texts_number % WORD_NUMBER_HINT_BY == 0:
                    mode = "feedback"
                    name = self.format_loc(LOC_LOG_WORD_NUMBER_HINT, name=self.record_dict[port].name, number=texts_number)
                    return True, False, (mode, name)
        #处理指令
        for key in ["log","记录"]:
            if not should_proc and msg_str.startswith(f".{key}"):
                should_proc = True
                msg_str = msg_str[1 + len(key):].strip()
        mode: str = ""
        name: str = ""
        if should_proc:
            for key in ["start","bew","开始","新建"]:
                if not mode and msg_str.startswith(key):
                    msg_str = msg_str[(len(key)):].strip()
                    mode = "start"
                    if len(msg_str) > 0:
                        name = msg_str
                    else:
                        name = str(meta.group_id) + " - " + "1"
            for key in ["list","列表"]:
                if not mode and msg_str.startswith(key):
                    msg_str = msg_str[(len(key)):].strip()
                    mode = "list"
            for key in ["download","下载"]:
                if not mode and msg_str.startswith(key):
                    msg_str = msg_str[(len(key)):].strip()
                    mode = "download"
                    name = msg_str
            for key in ["resume","on","继续"]:
                if not mode and msg_str.startswith(key):
                    msg_str = msg_str[(len(key)):].strip()
                    mode = "resume"
                    name = msg_str
            for key in ["regen","生成"]:
                if not mode and msg_str.startswith(key):
                    msg_str = msg_str[(len(key)):].strip()
                    mode = "regen"
                    name = msg_str
            for key in ["stop","off","暂停"]:
                if not mode and msg_str.startswith(key):
                    msg_str = msg_str[(len(key)):].strip()
                    mode = "stop"
                    name = msg_str
            for key in ["end","结束"]:
                if not mode and msg_str.startswith(key):
                    msg_str = msg_str[(len(key)):].strip()
                    mode = "end"
                    name = msg_str
            if not mode:
                if port in self.record_dict:
                    mode = "now"
                else:
                    if len(msg_str) > 0:
                        mode = "start"
                        name = msg_str
                    else:
                        mode = "help"
        return should_proc, should_pass, (mode, name)

    def process_msg(self, msg_str: str, meta: MessageMetaData, hint: Any) -> List[BotCommandBase]:
        port = GroupMessagePort(meta.group_id)
        mode = hint[0]
        name = hint[1]
        feedback: str = ""
        file_path = DATA_PATH + "\\LogData\\" + str(meta.group_id) + "\\" + name + ".txt"
        # 检查名称是否有效
        if name != "":
            for invalid_word in "\\/:?\"<>|":
                if invalid_word in name:
                    feedback = self.format_loc(LOC_LOG_INVALID_NAME)
                    return [BotSendMsgCommand(self.bot.account, feedback, [port])]
        # 处理
        if mode == "start":
            self.record_dict[port] = LogRecord(name,meta.group_id)
            self.bot.data_manager.set_data(DC_GROUPLOG, [meta.group_id,"log"],name)
            self.bot.data_manager.set_data(DC_GROUPLOG, [meta.group_id,"texts"],0)
            self.bot.data_manager.set_data(DC_GROUPLOG, [meta.group_id,"status"],True)
            feedback = self.format_loc(LOC_LOG_START, name=name)
        elif mode == "stop":
            if port in self.record_dict:
                if name == "":
                    name = self.record_dict[port].name
                if self.record_dict[port].mode == 0:
                    self.record_dict[port].mode = 1
                    self.bot.data_manager.set_data(DC_GROUPLOG, [meta.group_id,"status"],False)
                    feedback = self.format_loc(LOC_LOG_STOP, name=name)
                else:
                    feedback = self.format_loc(LOC_LOG_NOT_START, name=name)
            else:
                feedback = self.format_loc(LOC_LOG_NOT_START)
        elif mode == "resume":
            if port in self.record_dict:
                if name == "":
                    name = self.record_dict[port].name
                if self.record_dict[port].mode == 1:
                    self.record_dict[port].mode = 0
                    self.bot.data_manager.set_data(DC_GROUPLOG, [meta.group_id,"status"],True)
                    feedback = self.format_loc(LOC_LOG_RESUME, name=name)
                else:
                    feedback = self.format_loc(LOC_LOG_ALREADY_START, name=name)
            else:
                feedback = self.format_loc(LOC_LOG_NOT_FOUND, name=name)
        elif mode == "end":
            if port in self.record_dict:
                feedback = self.record_dict[port].generate_docx()
                self.bot.data_manager.set_data(DC_GROUPLOG, [meta.group_id,"log"],"")
                self.bot.data_manager.set_data(DC_GROUPLOG, [meta.group_id,"texts"],0)
                self.bot.data_manager.set_data(DC_GROUPLOG, [meta.group_id,"status"],False)
                feedback = self.format_loc(LOC_LOG_END, name=self.record_dict[port].name) + "\n" + feedback
                del self.record_dict[port]
            else:
                feedback = self.format_loc(LOC_LOG_NOT_START)
        elif mode == "list":
            dir_path = DATA_PATH + "\\LogData\\" + str(meta.group_id)
            if os.path.exists(dir_path):  # 遍历文件夹下所有文件
                try:
                    log_list: list = []
                    inner_paths = os.listdir(dir_path)
                    for inner_path in inner_paths:
                        if inner_path.endswith(".txt"):
                            log_list.append(os.path.basename(inner_path)[:-4])
                    feedback = self.format_loc(LOC_LOG_LIST, result=", ".join(log_list))
                except FileNotFoundError as e:  # 文件夹不存在
                    feedback = self.format_loc(LOC_LOG_NO_RESULT)
            else:
                feedback = self.format_loc(LOC_LOG_NO_RESULT)
        elif mode == "download":
            if name:
                if os.path.exists(file_path):
                    docx_path = DATA_PATH + "\\LogData\\" + str(meta.group_id) + "\\" + name + ".docx"
                    if os.path.exists(docx_path):
                        return [
                            BotSendFileCommand(self.bot.account, file_path, name + ".txt", [port]),
                            BotSendFileCommand(self.bot.account, docx_path, name + ".docx", [port])
                        ]
                    else:
                        return [BotSendFileCommand(self.bot.account, file_path, name + ".txt", [port])]
                else:
                    feedback = self.format_loc(LOC_LOG_NOT_FOUND, name=name)
            elif port in self.record_dict:
                feedback = self.record_dict[port].generate_docx(meta.user_id,self.format_loc(LOC_LOG_EMAIL, name=name))
                feedback = self.format_loc(LOC_LOG_REGEN, name=name) + "\n" + feedback
            else:
                feedback = self.format_loc(LOC_LOG_NOT_FOUND, name=name)
        elif mode == "regen":
            if name:
                if os.path.exists(file_path):
                    record = LogRecord(name,meta.group_id)
                    feedback = record.generate_docx(meta.user_id,self.format_loc(LOC_LOG_EMAIL, name=name))
                    del record
                    feedback = self.format_loc(LOC_LOG_REGEN, name=name) + "\n" + feedback
                else:
                    feedback = self.format_loc(LOC_LOG_NOT_FOUND, name=name)
            elif port in self.record_dict:
                feedback = self.record_dict[port].generate_docx(meta.user_id,self.format_loc(LOC_LOG_EMAIL, name=name))
                feedback = self.format_loc(LOC_LOG_REGEN, name=name) + "\n" + feedback
            else:
                feedback = self.format_loc(LOC_LOG_NOT_FOUND, name=name)
        elif mode == "help":
            feedback = self.get_help("log",meta)
        elif mode == "feedback":
            feedback = name
        else:
            if port in self.record_dict:
                feedback = self.format_loc(LOC_LOG_NOW, name=self.record_dict[port].name)
            else:
                feedback = self.format_loc(LOC_LOG_NOT_START)
        return [BotSendMsgCommand(self.bot.account, feedback, [port])]

    def get_help(self, keyword: str, meta: MessageMetaData) -> str:
        if keyword == "log":  # help后的接着的内容
            feedback: str = ".log 跑团记录系统"\
                "\n.log [跑团文件名称] 开启或重启跑团记录"\
                "\n.log off/stop 暂停跑团记录"\
                "\n.log on/resume 继续跑团记录"\
                "\n.log now 确认当前跑团记录开启情况"\
                "\n.log end 结束跑团记录并润色"\
                "\n.log regen [跑团文件名称] 重新润色跑团记录"\
                "\n润色完成的文件会与源文件一起发送至群文件或你的qq邮箱"\
                "\n请注意检查您的收件箱与垃圾箱"
            return feedback
        return ""

    def get_description(self) -> str:
        return ".log 跑团记录"  # help指令中返回的内容
